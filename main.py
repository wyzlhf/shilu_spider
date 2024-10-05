import time
from typing import List
from docx import Document
import requests
from bs4 import BeautifulSoup, ResultSet
from docx.shared import Pt
from pydantic import HttpUrl
from requests import Response


def get_soup_form_url(url: HttpUrl) -> BeautifulSoup:
    res: Response = requests.get(url)
    res.encoding = res.apparent_encoding
    response: str = res.text
    soup: BeautifulSoup = BeautifulSoup(response, 'html.parser')
    return soup


def create_document_with_name(name: str):
    document = Document()
    document.save(f'{name}.docx')


def write_h2_with_title(doc_name: str, title: str):
    doc = Document(f'{doc_name}.docx')
    doc.add_heading(title, level=2)
    doc.save(f'{doc_name}.docx')


def write_chapter_to_doc_with_content(content: str, doc_name: str):
    doc = Document(f'{doc_name}.docx')
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(28)
    run = para.add_run(content)
    run.font.name = '华文宋体'
    run.font.size = Pt(14)
    doc.save(f'{doc_name}.docx')


def get_all_book_from_home_page(url: HttpUrl) -> List[dict]:
    soup: BeautifulSoup = get_soup_form_url(url)
    book_div: ResultSet = soup.find_all(class_='col-xs-6 col-md-2')
    book_name_and_url_list: List[dict] = []
    for item in book_div:
        a_tag = item.find('a')
        book_url = a_tag.get('href')
        book_name: str = a_tag.text.strip()
        book_dict = {'book_name': book_name, 'book_url': book_url}
        book_name_and_url_list.append(book_dict)
    return book_name_and_url_list


def get_all_chapters_from_book_page(url: HttpUrl) -> List[dict]:
    soup: BeautifulSoup = get_soup_form_url(url)
    a_list: List = soup.find(id='booklist').find_all('a')
    chapter_struct: List[dict] = []
    for item in a_list:
        chapter_title: str = item.text
        chapter_href: HttpUrl = item.get('href')
        chapter_dict = {'chapter_title': chapter_title, 'chapter_href': chapter_href}
        chapter_struct.append(chapter_dict)
    return chapter_struct


def get_content_from_chapter_page(url: HttpUrl, doc_name: str):
    soup: BeautifulSoup = get_soup_form_url(url)
    content = soup.find(id='content').children
    chapter_text: str = ''
    for item in content:
        item_text = item.text.strip()
        if item_text != '':
            # write_chapter_to_doc_with_content(item_text, doc_name)
            chapter_text += item_text + '\n'
        else:
            continue
    write_chapter_to_doc_with_content(chapter_text, doc_name)


def main(qingshilu_url: HttpUrl):
    book_name_and_url_list = get_all_book_from_home_page(qingshilu_url)
    for book_item in book_name_and_url_list:
        book_name = book_item['book_name']
        book_url = book_item['book_url']
        create_document_with_name(book_name)
        print(
            f'=================={book_name}开始写入-{time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())}==================')
        chapter_struct = get_all_chapters_from_book_page(book_url)
        for chapter_item in chapter_struct:
            chapter_name: str = chapter_item['chapter_title']
            write_h2_with_title(book_name, chapter_name)
            chapter_href: HttpUrl = chapter_item['chapter_href']
            get_content_from_chapter_page(chapter_href, book_name)
            print(
                f'--------------------{chapter_name}写入完成-{time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())}----------------------------')
        print(
            f'=================={book_name}写入完成-{time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())}==========================')


if __name__ == '__main__':
    main('https://www.zhonghuadiancang.com/shudan/qingshilu/')
